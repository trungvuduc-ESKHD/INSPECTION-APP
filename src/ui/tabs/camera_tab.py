import streamlit as st
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import datetime
from src.core.supabase_client import upload_file, supabase

# ===================================================================
# HÀM TẠO BÁO CÁO WORD CHO MỘT SẢN PHẨM DUY NHẤT
# ===================================================================
def create_single_product_photo_report(product, product_images):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    product_name = product.get('name', 'Unknown Product')
    product_size = product.get('size', '')

    # --- Tiêu đề ---
    p = doc.add_paragraph()
    p.add_run('GENERAL PHOTO\n').bold = True
    p.add_run(f'TỔNG QUAN VỀ SẢN PHẨM/ PRODUCT OVERVIEW\n').bold = True
    p.add_run(f"{product_name} SIZE {product_size}").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Vẽ các bảng ảnh ---
    for category, images in product_images.items():
        if not images:
            continue

        p = doc.add_paragraph()
        p.add_run(category.replace('_', ' ').title()).bold = True # Chuyển "overview" thành "Overview"

        cols = 2
        rows = (len(images) + 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Thiết lập chiều rộng cột
        for col in table.columns:
            col.width = Inches(3.5)

        # Chèn ảnh vào các ô
        img_idx = 0
        for r in range(rows):
            for c in range(cols):
                if img_idx < len(images):
                    try:
                        cell = table.cell(r, c)
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(BytesIO(images[img_idx]), width=Inches(3.2))
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Error adding image to report: {e}")
                    img_idx += 1

        doc.add_paragraph()

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

def handle_new_photo():
    """
    Hàm này được gọi mỗi khi st.camera_input có ảnh mới.
    Nó sẽ lưu ảnh và tăng key của camera để làm mới widget.
    """
    # Lấy thông tin từ session state
    product_name = st.session_state.get('camera_current_product')
    category = st.session_state.get('camera_current_category')
    camera_key = st.session_state.get('camera_widget_key') # Key của camera_input

    if not product_name or not category or not camera_key:
        return

    # Lấy dữ liệu ảnh từ widget camera thông qua key của nó
    img_buffer = st.session_state[camera_key]
    
    if img_buffer:
        # Thêm ảnh vào danh sách
        st.session_state.camera_images[product_name][category].append(img_buffer.getvalue())
        # Tăng key của camera để "buộc" nó làm mới cho lần chụp tiếp theo
        st.session_state.camera_key_counter += 1

# ===================================================================
# HÀM RENDER CHÍNH CỦA TAB
# ===================================================================
def render_camera_tab():
    st.header("📸 Chụp ảnh báo cáo")

    # Khởi tạo các biến trong session state nếu chưa có
    if 'camera_images' not in st.session_state:
        st.session_state.camera_images = {}

    if not st.session_state.get('inspection_data') or not st.session_state.inspection_data.get('products'):
        st.warning("⚠️ Vui lòng thêm sản phẩm ở tab 'Thông tin/General' trước.")
        return

    # Giao diện chọn sản phẩm và danh mục ảnh
    product_names = [p['name'] for p in st.session_state.inspection_data['products']]
    selected_product_name = st.selectbox("1. Chọn sản phẩm", product_names, key="upload_product_selector")

    categories = {
        "Overview": "Tổng quan", "Checking_weight": "Kiểm tra khối lượng",
        "Checking_size": "Kiểm tra kích cỡ", "Checking_Brix": "Kiểm tra độ Brix",
        "Checking_Firmness": "Kiểm tra độ cứng", "Serious_defects": "Lỗi nghiêm trọng",
        "Major_defects": "Lỗi lớn", "Minor_defects": "Lỗi nhẹ",
        "Shattering_Berries": "Quả bị rụng"
    }
    selected_category = st.selectbox("2. Chọn danh mục ảnh", options=list(categories.keys()), format_func=lambda k: categories[k])

    # Khởi tạo dict lưu ảnh nếu chưa có
    if selected_product_name not in st.session_state.camera_images:
        st.session_state.camera_images[selected_product_name] = {cat: [] for cat in categories}

    st.markdown("---")
    st.subheader(f"📤 3. Tải ảnh hoặc chụp ảnh")

    st.info("📌 Bạn có thể chọn nhiều ảnh hoặc chụp trực tiếp bằng camera (nếu trình duyệt hỗ trợ).")
    uploaded_files = st.file_uploader(
        "Nhấn để chọn hoặc chụp ảnh (có thể chọn nhiều)",
        type=["jpg", "jpeg", "png"],
        accept_multiple_files=True,
        key="file_uploader"
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            img_bytes = uploaded_file.read()
            st.session_state.camera_images[selected_product_name][selected_category].append(img_bytes)
        st.success(f"✅ Đã thêm {len(uploaded_files)} ảnh cho mục '{categories[selected_category]}'")

    # Hiển thị các ảnh đã thêm
    st.markdown("---")
    st.subheader(f"🖼️ Các ảnh đã tải/chụp cho mục '{categories[selected_category]}'")

    current_images = st.session_state.camera_images[selected_product_name][selected_category]
    if not current_images:
        st.info("Chưa có ảnh nào.")
    else:
        cols = st.columns(4)
        for i, img in enumerate(current_images):
            with cols[i % 4]:
                st.image(img, use_container_width=True)
                if st.button(f"❌ Xóa ảnh {i+1}", key=f"del_{selected_product_name}_{selected_category}_{i}"):
                    del current_images[i]
                    st.rerun()


    # --- TẠO, UPLOAD VÀ TẢI BÁO CÁO (LOGIC MỚI) ---
    # ==========================================================
    st.markdown("---")
    st.header("Tạo và Tải báo cáo")
    st.info("Báo cáo sẽ được tạo và tự động lưu trữ online trên hệ thống.")

    # Lấy danh sách sản phẩm có ảnh
    products_with_images = []
    all_products = st.session_state.inspection_data.get('products', [])
    for product in all_products:
        if product['name'] in st.session_state.camera_images and any(st.session_state.camera_images[product['name']].values()):
            products_with_images.append(product)

    if not products_with_images:
        st.warning("Chưa có ảnh nào được chụp.")
    else:
        for product in products_with_images:
            product_name = product['name']
            
            # --- Tạo layout cho mỗi sản phẩm ---
            with st.container(border=True):
                st.subheader(f"Báo cáo cho: '{product_name}'")
                col1, col2 = st.columns([3, 1])

                with col1:
                    st.write("Nhấn nút để tạo, lưu trữ online và tải về file báo cáo ảnh.")
                
                with col2:
                    if st.button(f"Tạo & Lưu Báo cáo '{product_name}'", key=f"generate_{product_name}", use_container_width=True, type="primary"):
                        with st.spinner(f"Đang xử lý cho '{product_name}'..."):
                            
                            # 1. Tạo file .docx trong bộ nhớ
                            report_stream = create_single_product_photo_report(product, st.session_state.camera_images[product_name])
                            report_bytes = report_stream.getvalue()
                            
                            # 2. Chuẩn bị đường dẫn và tên file trên Supabase Storage
                            report_id = st.session_state.inspection_data['generalInfo']['report_id']
                            file_name_on_storage = f"{report_id}/{product_name.replace(' ', '_')}_Photo_Report.docx"
                            
                            # 3. Upload file lên Supabase
                            public_url = upload_file(
                                bucket_name="reportimages", # Tên bucket của bạn
                                file_path=file_name_on_storage,
                                file_body=report_bytes,
                                file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
                            )
                            
                            if public_url:
                                st.success(f"Đã lưu báo cáo lên hệ thống thành công!")
                                
                                # 4. (Tùy chọn) Cập nhật URL vào CSDL PostgreSQL
                                try:
                                    supabase.table('reports').update({'photo_report_url': public_url}).eq('id', report_id).execute()
                                except Exception as e:
                                    st.warning(f"Lưu báo cáo thành công nhưng không thể cập nhật link vào CSDL: {e}")

                                # 5. Cung cấp nút tải xuống cho người dùng
                                st.download_button(
                                    label=f"Tải lại file '{product_name}'",
                                    data=report_bytes,
                                    file_name=f"Photo_Report_{product_name.replace(' ', '_')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            else:
                                st.error("Có lỗi xảy ra khi tải báo cáo lên hệ thống.")
